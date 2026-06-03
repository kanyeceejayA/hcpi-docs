"""Build HCPI_Reported_Issues.docx from the triage data."""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date

# ---------- styling helpers ----------

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
ACCENT = RGBColor(0x2E, 0x86, 0xAB)
MUTED = RGBColor(0x55, 0x55, 0x55)
LIGHT_BG = "EEF3F8"
HEADER_BG = "1F3A5F"


def shade_cell(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell, color="BFBFBF"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), color)
        borders.append(b)
    tc_pr.append(borders)


def style_run(run, *, size=None, bold=None, color=None, font="Calibri"):
    run.font.name = font
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    if level == 0:
        style_run(run, size=26, bold=True, color=NAVY)
    elif level == 1:
        style_run(run, size=18, bold=True, color=NAVY)
    elif level == 2:
        style_run(run, size=14, bold=True, color=ACCENT)
    else:
        style_run(run, size=12, bold=True, color=MUTED)
    return p


def add_paragraph(doc, text, *, size=11, bold=False, color=None, italic=False, indent=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if indent is not None:
        p.paragraph_format.left_indent = indent
    run = p.add_run(text)
    style_run(run, size=size, bold=bold, color=color)
    run.italic = italic
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(item)
        style_run(run, size=11)


def add_kv_paragraph(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f"{label}: ")
    style_run(r1, size=11, bold=True, color=NAVY)
    r2 = p.add_run(value)
    style_run(r2, size=11)


def add_table(doc, headers, rows, col_widths_cm=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    if col_widths_cm:
        for col, width in zip(table.columns, col_widths_cm):
            for cell in col.cells:
                cell.width = Cm(width)

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(h)
        style_run(run, size=10.5, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        shade_cell(hdr_cells[i], HEADER_BG)
        set_cell_borders(hdr_cells[i])
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        bg = LIGHT_BG if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(row):
            cells[c_idx].text = ""
            p = cells[c_idx].paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(str(val))
            style_run(run, size=10)
            shade_cell(cells[c_idx], bg)
            set_cell_borders(cells[c_idx])
            cells[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        if col_widths_cm:
            for cell, width in zip(cells, col_widths_cm):
                cell.width = Cm(width)

    doc.add_paragraph()
    return table


def add_section(doc, title, intro, bullet_summary, headers, rows, col_widths_cm=None):
    add_heading(doc, title, level=1)
    add_paragraph(doc, intro, italic=True, color=MUTED)
    add_paragraph(doc, "At a glance", bold=True, color=ACCENT)
    add_bullets(doc, bullet_summary)
    add_paragraph(doc, "Details", bold=True, color=ACCENT)
    add_table(doc, headers, rows, col_widths_cm=col_widths_cm)


# ---------- document content ----------

doc = Document()

# margins
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

# default font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# --- title block ---
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run("HCPI System")
style_run(run, size=14, color=ACCENT, bold=True)

title_p2 = doc.add_paragraph()
title_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p2.add_run("Reported Issues Report")
style_run(run, size=28, bold=True, color=NAVY)

meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta_p.add_run(f"Prepared {date.today().strftime('%d %B %Y')}    ·    EAC HCPI Maintenance Team")
style_run(run, size=10.5, color=MUTED)
run.italic = True

doc.add_paragraph()

# --- executive summary ---
add_heading(doc, "Executive Summary", level=1)
add_paragraph(
    doc,
    "This document consolidates issues reported by the four national statistics offices "
    "(UBOS, KNBS, NBS, OCGS) and the EAC Secretariat against the HCPI deployment, together "
    "with a maintenance-team triage. Each issue is grouped by functional area, sized for "
    "effort, and assigned a target branch so the work can be planned, scheduled, and "
    "communicated back to the reporting offices.",
)
add_paragraph(
    doc,
    "Effort sizing uses three buckets: S (1–3 days), M (3–10 days), and L (open-ended "
    "or governance-dependent). “Shared” in the Branch column means the change lives on "
    "the 18.0 base and is merged forward into each country branch (ug_18, ke_18_mtnce, "
    "tz_18_mtnce, zar_18).",
)

# --- issues noticed ---
add_heading(doc, "Issues Noticed — Cross-Cutting Themes", level=1)
add_paragraph(
    doc,
    "Reading across the reports, four patterns emerge that recur in more than one country "
    "or more than one module:",
    italic=True,
    color=MUTED,
)
add_bullets(
    doc,
    [
        "Dashboard usability gaps. Several offices report the same UX complaints: no way to "
        "select a specific month / year, headline national figure not prominent, not all "
        "divisions visible. These are not country-specific bugs — they are shared "
        "dashboard limitations.",
        "Empty-data symptoms that look like bugs but are usually data state. “Sector figures "
        "show 0” (UBOS, KNBS) and “no national figures” (Uganda) most often trace back to "
        "the zero-price safety gate, missing dashboard_display flags, or the national index "
        "wizard not having run — not to broken code.",
        "Validation discoverability. Outliers and inliers are flagged in the database but "
        "have no dedicated report. The same problem applies to items that have been "
        "temporarily missing for several months, and to centrally-administered prices that "
        "move on their own cadence.",
        "Operational hygiene. Slowness (KNBS), advisory-lock leaks on deleted update history, "
        "and unbounded growth of log / tracking tables are symptoms of the same maintenance "
        "debt and should be addressed as one stability pass.",
    ],
)
add_paragraph(
    doc,
    "The remainder of this document breaks the full issue list into six functional groups, "
    "with a short bullet summary and a detail table for each.",
)

# --- groups ---

add_section(
    doc,
    "1. Dashboard & Reporting",
    intro=(
        "The dashboard is the most visible surface of the system and the source of the "
        "largest single cluster of issues. Most items here are usability or layout "
        "improvements; two are data-state investigations specific to Uganda and Kenya."
    ),
    bullet_summary=[
        "Add month / year filter so users can view any historical period.",
        "Lift the national CPI headline to the top of the page; raise the historical trend.",
        "Show all twelve COICOP divisions on screen, not just flagged ones.",
        "Add annual %% change and Core / Non-core cards alongside the existing series.",
        "Provide a public, login-free dashboard route for external audiences.",
        "Investigate Uganda dashboard showing no national figures.",
        "Investigate Kenya dashboard showing 0 for every sector.",
    ],
    headers=["#", "Issue", "Module / Area", "Effort", "Branch"],
    rows=[
        ("1.1", "Month / year dropdown filter", "hcpi_dashboard", "M", "shared"),
        ("1.2", "National CPI value at top; trend above the fold", "hcpi_dashboard", "M", "shared"),
        ("1.3", "Show all 12 Divisions on screen", "hcpi_dashboard + data", "S", "shared"),
        ("1.4", "Annual % change + Core / Non-core indicator cards", "hcpi_dashboard", "S", "shared"),
        ("1.5", "Public dashboard access (no login)", "new hcpi_dashboard_public", "M", "shared"),
        ("1.6", "Uganda / UBOS: no national figures on dashboard", "data investigation", "M", "ug_18"),
        ("1.7", "Kenya / KNBS: sector figures showing 0", "data investigation", "M", "ke_18_mtnce"),
    ],
    col_widths_cm=[1.0, 7.2, 4.0, 1.5, 3.3],
)

add_section(
    doc,
    "2. Data Collection & Validation",
    intro=(
        "Field workflow and quality control. Most items are additive (new reports, smarter "
        "checks) and ship from the shared 18.0 branch. The Kenya progress-bar bug and the "
        "Tanzania offline-confirmation request are the only country-flavoured items."
    ),
    bullet_summary=[
        "Build a dedicated report of outlier and inlier flags raised by validation.",
        "Surface items that have been temporarily missing for three months or more.",
        "Add a Data Collection Analytics view: counts by state, by outlet / centre / national.",
        "Fix the KNBS progress bar dropping from 100%% to 67%% after validation.",
        "Extend the Tukey outlier check into three severity tiers (24 / 12 / 6 months).",
        "Strengthen the inlier check to detect identical month-on-month change patterns.",
        "Flag changes to centrally-administered prices for separate review.",
        "Ensure collected data reliably reaches the permanent observations table.",
        "Confirm and document offline behaviour of the collection tablet (NBS request).",
    ],
    headers=["#", "Issue", "Module / Area", "Effort", "Branch"],
    rows=[
        ("2.1", "Report of flagged outliers and inliers", "hcpi_data_collection", "S", "shared"),
        ("2.2", "Indicator + report for items missing 3+ months", "hcpi_outlet", "S", "shared"),
        ("2.3", "Data Collection Analytics by state and geography", "hcpi_data_collection", "M", "shared"),
        ("2.4", "KNBS progress bar: 100% then drops to 67%", "hcpi_data_collection", "M", "shared"),
        ("2.5", "Three-tier outliers (24 / 12 / 6 mo)", "hcpi_data_collection (run_turkey)", "M", "shared"),
        ("2.6", "Smarter inliers: same-change pattern check", "hcpi_data_collection (validate_inliers)", "M", "shared"),
        ("2.7", "Flag changes in centrally-administered prices", "hcpi_item + small dashboard", "M", "shared"),
        ("2.8", "Collected data must reach observations", "hcpi_data_collection (action_approve)", "M", "shared"),
        ("2.9", "NBS: confirm offline tablet behaviour", "test + docs", "S", "docs"),
    ],
    col_widths_cm=[1.0, 6.5, 5.2, 1.5, 2.8],
)

add_section(
    doc,
    "3. Computations & Index Logic",
    intro=(
        "Changes to how the index itself is produced. The formal review is already drafted "
        "in HCPI_COMPUTATION_FLOW.md — the remaining work is exposing computation detail to "
        "statisticians, giving operators more control over recomputation, and a methodology-"
        "dependent provisional-index workflow."
    ),
    bullet_summary=[
        "Publish the formal review of how computations work (already drafted).",
        "Add a “show the math” view for individual EA index values (KNBS request).",
        "Add Period and Division filters to all index-update wizards.",
        "Provisional index that carries forward last month’s values (methodology decision needed).",
        "Fix advisory-lock leak when update history is deleted, blocking restart.",
        "Clarify KNBS request on “standard and non-standard of weights.”",
    ],
    headers=["#", "Issue", "Module / Area", "Effort", "Branch"],
    rows=[
        ("3.1", "Formal review of computations process", "HCPI_COMPUTATION_FLOW.md (drafted)", "S", "docs"),
        ("3.2", "Show-the-math / in-built calculator view", "hcpi_index", "M", "shared"),
        ("3.3", "Period + Division filters on index update wizards", "hcpi_index", "M", "shared"),
        ("3.4", "Provisional Index — carry-forward methodology", "hcpi_index (new wizard)", "L", "shared"),
        ("3.5", "Deleted history leaves advisory lock", "hcpi_index queue-job", "M", "shared"),
        ("3.6", "KNBS: standard vs non-standard of weights", "needs clarification", "TBD", "TBD"),
    ],
    col_widths_cm=[1.0, 6.5, 5.2, 1.5, 2.8],
)

add_section(
    doc,
    "4. Performance & Infrastructure",
    intro=(
        "Reported by KNBS in particular but felt across instances as data accumulates. These "
        "two items are closely related and should be treated as a single stability pass."
    ),
    bullet_summary=[
        "Profile and address slowness as user count and historical data grow.",
        "Audit mail.message / ir_logging growth and add retention; tune tracking=True usage.",
    ],
    headers=["#", "Issue", "Module / Area", "Effort", "Branch"],
    rows=[
        ("4.1", "System slow with more users / data (KNBS)", "profiling first", "L", "shared + per-instance config"),
        ("4.2", "Large log tables inflating DB and slowing queries", "tracking audit + retention cron", "M", "shared"),
    ],
    col_widths_cm=[1.0, 6.8, 5.0, 1.5, 2.7],
)

add_section(
    doc,
    "5. Cross-Country & Governance",
    intro=(
        "Items that span multiple countries, require coordination with EAC, or are "
        "administrative rather than engineering tasks."
    ),
    bullet_summary=[
        "Review the EAC Secretariat sync — cron schedules, integration users, API keys.",
        "Standardise item codes across the four countries so cross-country reporting is possible.",
        "Publish a documentation page enumerating the XML-RPC endpoints the mobile app uses.",
        "Transfer the mobile app from Kola’s Google Play account to the EAC-controlled account.",
    ],
    headers=["#", "Issue", "Module / Area", "Effort", "Branch"],
    rows=[
        ("5.1", "Review syncing to Secretariat Dashboard (EAC)", "hcpi_secretariat_hub + per-country users", "M", "secretariat"),
        ("5.2", "Standardise item codes across countries", "hcpi_item + EAC-wide mapping", "L", "shared + data per country"),
        ("5.3", "Document XML / RPC endpoints", "docs page", "S", "docs"),
        ("5.4", "Transfer mobile app to EAC Google Play", "admin / account transfer", "S", "N/A"),
    ],
    col_widths_cm=[1.0, 6.5, 5.5, 1.5, 2.5],
)

add_section(
    doc,
    "6. Mobile App (Separate Repository)",
    intro=(
        "The Flutter mobile app is in a separate code repository — server-side changes here "
        "are only needed if new fields must be received on submission."
    ),
    bullet_summary=[
        "Investigate why GPS coordinates are not captured during data collection (NBS).",
    ],
    headers=["#", "Issue", "Module / Area", "Effort", "Branch"],
    rows=[
        ("6.1", "NBS: GPS not captured during data collection", "Flutter app", "L", "mobile app repo"),
    ],
    col_widths_cm=[1.0, 7.5, 4.7, 1.5, 2.3],
)

add_section(
    doc,
    "7. Already Resolved",
    intro="Items reported but already closed at the time of writing.",
    bullet_summary=[
        "OCGS: unable to update EA Indices — resolved.",
    ],
    headers=["#", "Issue", "Status"],
    rows=[
        ("7.1", "OCGS: unable to update EA Indices", "Resolved"),
    ],
    col_widths_cm=[1.0, 10.5, 5.5],
)

# --- sequencing ---
add_heading(doc, "Recommended Sequencing", level=1)
add_paragraph(
    doc,
    "If picked up in order, the items below answer the largest share of the reported "
    "complaints with the lowest implementation risk.",
)

add_paragraph(doc, "Sprint 1 — Dashboard refresh (one shared PR)", bold=True, color=ACCENT)
add_bullets(doc, ["1.1, 1.2, 1.3, 1.4, 2.1, 2.2, 2.3"])

add_paragraph(doc, "Sprint 2 — Stability pass", bold=True, color=ACCENT)
add_bullets(doc, ["2.4 (progress bug), 4.1 + 4.2 (slowness and log retention together), 3.5 (lock cleanup)"])

add_paragraph(doc, "Sprint 3 — Validation enhancements (methodology sign-off needed)", bold=True, color=ACCENT)
add_bullets(doc, ["2.5, 2.6, 2.7, 3.2, 3.3"])

add_paragraph(doc, "Investigation backlog (do before further sizing)", bold=True, color=ACCENT)
add_bullets(doc, [
    "1.6, 1.7 — require database snapshots from Uganda and Kenya.",
    "5.1 — audit Secretariat sync state with each country team.",
    "3.6 — clarify KNBS request with the reporter.",
])

add_paragraph(doc, "Governance and non-engineering", bold=True, color=ACCENT)
add_bullets(doc, [
    "3.1 publish review, 3.4 methodology sign-off, 5.2 cross-country code agreement, 5.4 Google Play transfer.",
])

# --- footer note ---
doc.add_paragraph()
add_paragraph(
    doc,
    "Source: reported-issues.md. Detailed per-issue analysis (root cause, files touched, "
    "rationale): see issue-triage.md.",
    size=9.5,
    italic=True,
    color=MUTED,
)

out_path = r"c:\code\hcpi-docs\HCPI_Reported_Issues.docx"
doc.save(out_path)
print(f"Wrote {out_path}")
